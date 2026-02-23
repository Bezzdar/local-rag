"""Сервис парсинга и нормализации документов.

Поддерживает несколько методов чанкинга (General, Context Enrichment,
Hierarchy, PCR, Symbol) и возвращает унифицированную структуру чанков,
чтобы слой хранения/поиска не зависел от выбранного алгоритма.
"""

# --- Imports ---
from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Optional
from uuid import uuid4

from langdetect import detect

from ..config import CHUNKS_DIR

try:
    import tiktoken
except Exception:  # noqa: BLE001
    tiktoken = None


# --- Основные блоки ---
class UnsupportedFormatError(Exception):
    pass


class ParseError(Exception):
    pass


class ChunkType(Enum):
    TEXT = "text"
    TABLE = "table"
    FORMULA = "formula"
    HEADER = "header"
    CAPTION = "caption"


# Supported chunking methods and document types
CHUNKING_METHODS = ["general", "context_enrichment", "hierarchy", "pcr", "symbol"]
DOC_TYPES = ["technical_manual", "gost", "api_docs", "markdown"]


@dataclass
class ParserConfig:
    """Глобальные настройки парсинга и чанкинга.

    Значения выступают дефолтами и могут переопределяться для конкретного
    источника через ``individual_config``.
    """
    chunk_size: int = 512
    chunk_overlap: int = 64
    min_chunk_size: int = 50
    ocr_enabled: bool = True
    ocr_language: str = "rus+eng"
    # Chunking method selection
    chunking_method: str = "general"
    # Context Enrichment params
    context_window: int = 128
    use_llm_summary: bool = False
    # Hierarchy params
    doc_type: str = "technical_manual"
    # PCR (Parent-Child Retrieval) params
    parent_chunk_size: int = 1024
    child_chunk_size: int = 128
    # Symbol separator params
    symbol_separator: str = "---chunk---"


@dataclass
class ParsedChunk:
    """Нормализованная модель чанка для БД, поиска и ответа LLM."""
    text: str
    chunk_type: ChunkType
    chunk_index: int
    page_number: Optional[int]
    section_header: Optional[str]
    parent_header: Optional[str]
    prev_chunk_tail: Optional[str]
    next_chunk_head: Optional[str]
    doc_id: str
    source_filename: str
    # Optional fields for advanced chunking methods
    embedding_text: Optional[str] = None   # Text used for embedding (differs from display text in CE/PCR)
    parent_chunk_id: Optional[str] = None  # For PCR: child chunks reference their parent


@dataclass
class DocumentMetadata:
    """Метаданные документа и конфигурации, с которой он был распарсен."""
    doc_id: str
    notebook_id: str
    filename: str
    filepath: str
    file_hash: str
    file_size_bytes: int
    title: Optional[str]
    authors: Optional[list[str]]
    year: Optional[int]
    source: Optional[str]
    total_pages: Optional[int]
    total_chunks: int
    language: str
    parser_version: str
    parsed_at: str
    tags: list[str] = field(default_factory=list)
    user_notes: Optional[str] = None
    is_enabled: bool = True
    individual_config: dict[str, int | bool | str | None] = field(
        default_factory=lambda: {
            "chunk_size": None,
            "chunk_overlap": None,
            "ocr_enabled": None,
            "ocr_language": None,
            "chunking_method": None,
            "context_window": None,
            "use_llm_summary": None,
            "doc_type": None,
            "parent_chunk_size": None,
            "child_chunk_size": None,
            "symbol_separator": None,
        }
    )
    chunking_method: str = "general"


def _tokenize(text: str) -> list[str]:
    """Простейшая токенизация по пробелам (fallback-режим)."""
    return text.split()


def _token_count(text: str) -> int:
    """Подсчет токенов через tiktoken, либо приближенная оценка длины."""
    if not text.strip():
        return 0
    if tiktoken is not None:
        try:
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except Exception:  # noqa: BLE001
            pass
    return max(1, int(len(_tokenize(text)) * 1.3))


def _sort_pdf_lines_multicolumn(lines: list[tuple[float, float, str, float]]) -> list[tuple[float, float, str, float]]:
    """Пытается восстановить порядок чтения для двухколоночного PDF."""
    if len(lines) < 3:
        return sorted(lines, key=lambda item: (item[0], item[1]))

    sorted_by_x = sorted(lines, key=lambda item: item[1])
    xs = [item[1] for item in sorted_by_x]
    gaps = [xs[idx + 1] - xs[idx] for idx in range(len(xs) - 1)]
    split_gap = max(gaps) if gaps else 0
    if split_gap < 80:
        return sorted(lines, key=lambda item: (item[0], item[1]))

    split_idx = gaps.index(split_gap) + 1
    split_x = (xs[split_idx - 1] + xs[split_idx]) / 2

    left = [line for line in lines if line[1] <= split_x]
    right = [line for line in lines if line[1] > split_x]
    left.sort(key=lambda item: item[0])
    right.sort(key=lambda item: item[0])
    return left + right


class DocumentParser:
    def __init__(self, config: ParserConfig):
        self.config = config

    def parse(
        self,
        filepath: str,
        notebook_id: str,
        metadata_override: Optional[dict] = None,
    ) -> tuple[DocumentMetadata, list[ParsedChunk]]:
        """Полный пайплайн парсинга документа.

        Этапы:
        1) извлечение сырых блоков из файла;
        2) разбиение на чанки выбранным методом;
        3) сбор метаданных;
        4) сохранение результата в ``CHUNKS_DIR``.
        """
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(path)

        # Внешний слой может передать фиксированный doc_id и индивидуальные настройки.
        metadata_override = metadata_override or {}
        suffix = path.suffix.lower()
        # На этом этапе получаем унифицированные блоки, независимо от формата файла.
        blocks, total_pages = self._extract_blocks(path, suffix)

        doc_id = str(metadata_override.get("doc_id") or uuid4())
        # Далее блоки маршрутизируются в выбранный алгоритм чанкинга.
        chunks = self._chunk_blocks_dispatch(blocks, doc_id=doc_id, source_filename=path.name)

        metadata = DocumentMetadata(
            doc_id=doc_id,
            notebook_id=notebook_id,
            filename=path.name,
            filepath=str(path),
            file_hash=hashlib.sha256(path.read_bytes()).hexdigest(),
            file_size_bytes=path.stat().st_size,
            title=metadata_override.get("title"),
            authors=metadata_override.get("authors"),
            year=metadata_override.get("year"),
            source=metadata_override.get("source"),
            total_pages=total_pages,
            total_chunks=len(chunks),
            language=self.detect_language("\n".join(block["text"] for block in blocks)[:1000]),
            parser_version="1.1.0",
            parsed_at=datetime.now(timezone.utc).isoformat(),
            individual_config=metadata_override.get("individual_config")
            or {
                "chunk_size": None,
                "chunk_overlap": None,
                "ocr_enabled": None,
                "ocr_language": None,
                "chunking_method": None,
                "context_window": None,
                "use_llm_summary": None,
                "doc_type": None,
                "parent_chunk_size": None,
                "child_chunk_size": None,
                "symbol_separator": None,
            },
            is_enabled=bool(metadata_override.get("is_enabled", True)),
            chunking_method=self.config.chunking_method,
        )
        # Сохраняем промежуточный JSON, который затем потребляет embedding_service.
        self.save_parsing_result(notebook_id, metadata, chunks)
        return metadata, chunks

    def _extract_blocks(self, path: Path, suffix: str) -> tuple[list[dict], Optional[int]]:
        """Выбирает extraction-стратегию по расширению файла.

        Возвращает список унифицированных блоков и общее число страниц (если известно).
        """
        # Простой текст и markdown обрабатываем без внешних библиотек.
        if suffix in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            return self._text_to_structured_blocks(text, page_number=1), 1
        # DOCX: парсим параграфы/стили и таблицы.
        if suffix == ".docx":
            return self._extract_docx(path)
        # PDF: сначала text-layer, затем OCR fallback.
        if suffix == ".pdf":
            return self._extract_pdf(path)
        # XLSX пока обработан заглушкой (schema-ready, без полноценного extraction).
        if suffix == ".xlsx":
            return ([{
                "text": f"Table content placeholder for {path.name}",
                "chunk_type": ChunkType.TABLE,
                "page_number": 1,
                "section_header": None,
                "parent_header": None,
            }], 1)
        if suffix in {".html", ".epub"}:
            raise UnsupportedFormatError(f"Format planned but not implemented yet: {suffix}")
        raise UnsupportedFormatError(f"Unsupported format: {suffix}")

    def _text_to_structured_blocks(self, text: str, page_number: int) -> list[dict]:
        """Нормализует plain text в блоки HEADER/TEXT для общего пайплайна."""
        blocks: list[dict] = []
        current_header: Optional[str] = None
        # Идем построчно: заголовки помечаем отдельно, чтобы не терять структуру документа.
        for line in [ln.strip() for ln in text.splitlines() if ln.strip()]:
            is_header = bool(re.match(r"^(#{1,6}\s+.+|\d+(?:\.\d+)*\s+.+)$", line))
            if is_header:
                current_header = re.sub(r"^#{1,6}\s*", "", line)
                blocks.append(
                    {
                        "text": current_header,
                        "chunk_type": ChunkType.HEADER,
                        "page_number": page_number,
                        "section_header": current_header,
                        "parent_header": None,
                    }
                )
                continue
            blocks.append(
                {
                    "text": line,
                    "chunk_type": ChunkType.TEXT,
                    "page_number": page_number,
                    "section_header": current_header,
                    "parent_header": None,
                }
            )
        return blocks

    def _extract_docx(self, path: Path) -> tuple[list[dict], Optional[int]]:
        """Извлекает DOCX в список блоков; таблицы переводит в markdown-вид."""
        try:
            from docx import Document
            doc = Document(path)
        except Exception:
            return ([{
                "text": f"Extracted content from {path.name}",
                "chunk_type": ChunkType.TEXT,
                "page_number": None,
                "section_header": None,
                "parent_header": None,
            }], None)
        blocks: list[dict] = []
        current_header: Optional[str] = None
        # Параграфы DOCX конвертируем в блоки с учетом стилей (heading/list/plain).
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style is not None else "").lower()
            if "heading" in style:
                current_header = text
                blocks.append(
                    {
                        "text": text,
                        "chunk_type": ChunkType.HEADER,
                        "page_number": None,
                        "section_header": current_header,
                        "parent_header": None,
                    }
                )
                continue
            if "list" in style:
                marker = "- "
                blocks.append(
                    {
                        "text": f"{marker}{text}",
                        "chunk_type": ChunkType.TEXT,
                        "page_number": None,
                        "section_header": current_header,
                        "parent_header": None,
                    }
                )
                continue
            blocks.append(
                {
                    "text": text,
                    "chunk_type": ChunkType.TEXT,
                    "page_number": None,
                    "section_header": current_header,
                    "parent_header": None,
                }
            )

        # Таблицы приводим к markdown-представлению, чтобы их можно было чанковать как текст.
        for table in doc.tables:
            rows = [[cell.text.strip().replace("|", "\\|") for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = rows[0]
            divider = ["---"] * len(header)
            md_lines = [f"| {' | '.join(header)} |", f"| {' | '.join(divider)} |"]
            for row in rows[1:]:
                md_lines.append(f"| {' | '.join(row)} |")
            blocks.append(
                {
                    "text": "\n".join(md_lines),
                    "chunk_type": ChunkType.TABLE,
                    "page_number": None,
                    "section_header": current_header,
                    "parent_header": None,
                }
            )

        return blocks, None

    def _extract_pdf(self, path: Path) -> tuple[list[dict], Optional[int]]:
        """Извлекает PDF из text-layer или переключается на OCR при необходимости."""
        try:
            import fitz
        except Exception:
            return ([{
                "text": f"Extracted content from {path.name}",
                "chunk_type": ChunkType.TEXT,
                "page_number": 1,
                "section_header": None,
                "parent_header": None,
            }], 1)

        blocks: list[dict] = []
        section_header: Optional[str] = None
        try:
            doc_ctx = fitz.open(path)
        except Exception:
            return ([{
                "text": f"Extracted content from {path.name}",
                "chunk_type": ChunkType.TEXT,
                "page_number": 1,
                "section_header": None,
                "parent_header": None,
            }], 1)

        with doc_ctx as doc:
            total_pages = doc.page_count
            # Проверяем наличие текстового слоя: это ключевая развилка text-layer vs OCR.
            text_layer_present = any(page.get_text("text").strip() for page in doc)

            if not text_layer_present:
                if not self.config.ocr_enabled:
                    raise ParseError("Scanned PDF detected but OCR is disabled")
                return self._extract_pdf_ocr(path), total_pages

            # Постранично строим список строк с координатами и размером шрифта.
            for page_index in range(total_pages):
                page = doc.load_page(page_index)
                data = page.get_text("dict")
                lines: list[tuple[float, float, str, float]] = []
                for block in data.get("blocks", []):
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        if not spans:
                            continue
                        text = "".join(span.get("text", "") for span in spans).strip()
                        if not text:
                            continue
                        size = max(span.get("size", 11.0) for span in spans)
                        x0, y0, *_ = line.get("bbox", [0.0, 0.0, 0.0, 0.0])
                        lines.append((y0, x0, text, size))

                # Переупорядочиваем строки для двухколоночных макетов.
                lines = _sort_pdf_lines_multicolumn(lines)
                base_font = min((item[3] for item in lines), default=11.0)
                # Эвристика: увеличенный шрифт считаем заголовком, остальное — текстом.
                for _, _, text, size in lines:
                    if re.match(r"^\d+$", text):
                        continue
                    if size >= base_font + 1.5:
                        section_header = text
                        blocks.append(
                            {
                                "text": text,
                                "chunk_type": ChunkType.HEADER,
                                "page_number": page_index + 1,
                                "section_header": section_header,
                                "parent_header": None,
                            }
                        )
                    else:
                        blocks.append(
                            {
                                "text": text,
                                "chunk_type": ChunkType.TEXT,
                                "page_number": page_index + 1,
                                "section_header": section_header,
                                "parent_header": None,
                            }
                        )

                for image_idx, _image in enumerate(page.get_images(full=True), start=1):
                    blocks.append(
                        {
                            "text": f"[FORMULA_IMAGE: page_{page_index + 1}_formula_{image_idx}]",
                            "chunk_type": ChunkType.FORMULA,
                            "page_number": page_index + 1,
                            "section_header": section_header,
                            "parent_header": None,
                        }
                    )

        return blocks, total_pages


    def _preprocess_ocr_image(self, img, cv2):
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.fastNlMeansDenoising(gray)
        coords = cv2.findNonZero(255 - gray)
        if coords is not None:
            rect = cv2.minAreaRect(coords)
            angle = rect[-1]
            if angle < -45:
                angle = 90 + angle
            if abs(angle) > 0.3:
                h, w = gray.shape[:2]
                matrix = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
                gray = cv2.warpAffine(gray, matrix, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        return cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    def _extract_pdf_ocr(self, path: Path) -> list[dict]:
        """OCR-ветка для сканов: page -> image -> preprocess -> text blocks."""
        try:
            import cv2
            import fitz
            import numpy as np
            import pytesseract
        except Exception as exc:  # noqa: BLE001
            raise ParseError("OCR parsing requires opencv-python, pytesseract and PyMuPDF") from exc

        blocks: list[dict] = []
        try:
            doc_ctx = fitz.open(path)
        except Exception:
            return [{
                "text": f"Extracted content from {path.name}",
                "chunk_type": ChunkType.TEXT,
                "page_number": 1,
                "section_header": None,
                "parent_header": None,
            }]

        with doc_ctx as doc:
            # OCR выполняется постранично: растеризация -> preprocessing -> Tesseract.
            for page_idx in range(doc.page_count):
                page = doc.load_page(page_idx)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                if pix.n == 4:
                    img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
                preprocessed = self._preprocess_ocr_image(img, cv2)
                text = pytesseract.image_to_string(preprocessed, lang=self.config.ocr_language).strip()
                if text:
                    blocks.extend(self._text_to_structured_blocks(text, page_idx + 1))
        return blocks

    # --- Chunking dispatch ---

    def _chunk_blocks_dispatch(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """Маршрутизатор методов чанкинга по ``self.config.chunking_method``."""
        method = self.config.chunking_method
        if method == "context_enrichment":
            return self._chunk_method_context_enrichment(blocks, doc_id, source_filename)
        if method == "hierarchy":
            return self._chunk_method_hierarchy(blocks, doc_id, source_filename)
        if method == "pcr":
            return self._chunk_method_pcr(blocks, doc_id, source_filename)
        if method == "symbol":
            return self._chunk_method_symbol(blocks, doc_id, source_filename)
        # Default: general
        return self._chunk_blocks(blocks, doc_id=doc_id, source_filename=source_filename)

    # --- Method 1: General (existing) ---

    def _chunk_blocks(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """General-метод: fixed-size чанки + overlap соседних фрагментов."""
        chunks: list[ParsedChunk] = []
        pending_header: Optional[dict] = None

        # Проходим по потоку блоков и собираем единый список чанков.
        for block in blocks:
            block_type: ChunkType = block["chunk_type"]
            if block_type == ChunkType.HEADER:
                pending_header = block
                continue

            text = block["text"].strip()
            if not text:
                continue

            if pending_header is not None:
                text = f"{pending_header['text']}\n{text}"
                block["section_header"] = pending_header["text"]
                pending_header = None

            if block_type == ChunkType.TABLE:
                chunks.extend(self._chunk_table_block(text, block, doc_id, source_filename, len(chunks)))
                continue

            chunks.extend(self._chunk_text_block(text, block, doc_id, source_filename, len(chunks), block_type))

        # apply overlap metadata only
        # На финальном проходе записываем контекст соседей для retriever/LLM.
        overlap = max(0, self.config.chunk_overlap)
        for idx, chunk in enumerate(chunks):
            if idx > 0:
                chunk.prev_chunk_tail = " ".join(_tokenize(chunks[idx - 1].text)[-overlap:]) if overlap else None
            if idx < len(chunks) - 1:
                chunk.next_chunk_head = " ".join(_tokenize(chunks[idx + 1].text)[:overlap]) if overlap else None

        for idx, chunk in enumerate(chunks):
            chunk.chunk_index = idx
        return chunks

    # --- Низкоуровневые помощники чанкинга ---

    def _chunk_text_block(
        self,
        text: str,
        block: dict,
        doc_id: str,
        source_filename: str,
        start_index: int,
        chunk_type: ChunkType,
    ) -> list[ParsedChunk]:
        """Нарезает текст токен-окнами; при слишком коротком хвосте объединяет окна."""
        tokens = _tokenize(text)
        if not tokens:
            return []

        out: list[ParsedChunk] = []
        step = max(1, self.config.chunk_size)
        for offset in range(0, len(tokens), step):
            part_tokens = tokens[offset : offset + step]
            part_text = " ".join(part_tokens).strip()
            if _token_count(part_text) < self.config.min_chunk_size and offset + step < len(tokens):
                part_tokens = tokens[offset : offset + (step * 2)]
                part_text = " ".join(part_tokens).strip()
            out.append(
                ParsedChunk(
                    text=part_text,
                    chunk_type=chunk_type,
                    chunk_index=start_index + len(out),
                    page_number=block.get("page_number"),
                    section_header=block.get("section_header"),
                    parent_header=block.get("parent_header"),
                    prev_chunk_tail=None,
                    next_chunk_head=None,
                    doc_id=doc_id,
                    source_filename=source_filename,
                )
            )
        return out

    def _chunk_table_block(self, text: str, block: dict, doc_id: str, source_filename: str, start_index: int) -> list[ParsedChunk]:
        """Нарезает таблицу по строкам, дублируя заголовок в каждом куске."""
        lines = [ln for ln in text.splitlines() if ln.strip()]
        if len(lines) <= 2:
            return [
                ParsedChunk(
                    text=text,
                    chunk_type=ChunkType.TABLE,
                    chunk_index=start_index,
                    page_number=block.get("page_number"),
                    section_header=block.get("section_header"),
                    parent_header=block.get("parent_header"),
                    prev_chunk_tail=None,
                    next_chunk_head=None,
                    doc_id=doc_id,
                    source_filename=source_filename,
                )
            ]

        header = lines[:2]
        body = lines[2:]
        chunks: list[ParsedChunk] = []
        current_rows: list[str] = []
        for row in body:
            candidate = "\n".join(header + current_rows + [row])
            if _token_count(candidate) > self.config.chunk_size and current_rows:
                chunks.append(
                    ParsedChunk(
                        text="\n".join(header + current_rows),
                        chunk_type=ChunkType.TABLE,
                        chunk_index=start_index + len(chunks),
                        page_number=block.get("page_number"),
                        section_header=block.get("section_header"),
                        parent_header=block.get("parent_header"),
                        prev_chunk_tail=None,
                        next_chunk_head=None,
                        doc_id=doc_id,
                        source_filename=source_filename,
                    )
                )
                current_rows = [row]
                continue
            current_rows.append(row)

        if current_rows:
            chunks.append(
                ParsedChunk(
                    text="\n".join(header + current_rows),
                    chunk_type=ChunkType.TABLE,
                    chunk_index=start_index + len(chunks),
                    page_number=block.get("page_number"),
                    section_header=block.get("section_header"),
                    parent_header=block.get("parent_header"),
                    prev_chunk_tail=None,
                    next_chunk_head=None,
                    doc_id=doc_id,
                    source_filename=source_filename,
                )
            )
        return chunks

    # --- Method 2: Context Enrichment ---

    def _chunk_method_context_enrichment(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """Context Enrichment: каждый чанк получает контекстную обёртку из соседних фрагментов."""
        # Step 1: базовое разбиение (как General).
        chunks = self._chunk_blocks(blocks, doc_id=doc_id, source_filename=source_filename)

        # Step 2: формируем embedding_text из текущего чанка и контекста соседей.
        cw = max(0, self.config.context_window)
        for i, chunk in enumerate(chunks):
            prev_ctx = chunks[i - 1].text[-cw:] if i > 0 and cw > 0 else ""
            next_ctx = chunks[i + 1].text[:cw] if i < len(chunks) - 1 and cw > 0 else ""
            parts = [p for p in [prev_ctx, chunk.text, next_ctx] if p]
            chunk.embedding_text = " ".join(parts) if len(parts) > 1 else None

        return chunks

    # --- Method 3: Hierarchy ---

    # Patterns for different document types
    _HIERARCHY_PATTERNS = {
        "gost": [
            (1, re.compile(r"^\d+\.\s+[А-ЯA-Z\w]")),
            (2, re.compile(r"^\d+\.\d+\.\s+")),
            (3, re.compile(r"^\d+\.\d+\.\d+\.\s+")),
        ],
        "technical_manual": [
            (1, re.compile(r"^(Глава|Chapter|РАЗДЕЛ|SECTION)\s+\d+", re.IGNORECASE)),
            (2, re.compile(r"^\d+\.\d+\s+[А-ЯA-Z\w]")),
            (3, re.compile(r"^\d+\.\d+\.\d+\s+")),
        ],
        "api_docs": [
            (1, re.compile(r"^#{1,2}\s+")),
            (2, re.compile(r"^#{3}\s+")),
            (3, re.compile(r"^#{4,}\s+")),
        ],
        "markdown": [
            (1, re.compile(r"^#\s+")),
            (2, re.compile(r"^##\s+")),
            (3, re.compile(r"^###\s+")),
            (4, re.compile(r"^#{4,}\s+")),
        ],
    }

    def _detect_header_level(self, text: str, patterns: list[tuple[int, re.Pattern]]) -> Optional[int]:
        for level, pattern in patterns:
            if pattern.match(text):
                return level
        return None

    def _build_breadcrumb(self, hierarchy: dict[int, str]) -> str:
        parts = [hierarchy[level] for level in sorted(hierarchy.keys()) if hierarchy.get(level)]
        return " > ".join(parts) if parts else ""

    def _chunk_method_hierarchy(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """Hierarchy: разбиение по структурным маркерам документа с breadcrumb."""
        # Берем только заранее известные паттерны, иначе откатываемся к technical_manual.
        doc_type = self.config.doc_type if self.config.doc_type in self._HIERARCHY_PATTERNS else "technical_manual"
        patterns = self._HIERARCHY_PATTERNS[doc_type]

        chunks: list[ParsedChunk] = []
        hierarchy: dict[int, str] = {}  # level -> header text
        current_content_blocks: list[dict] = []

        # flush фиксирует накопленный раздел в один или несколько финальных чанков.
        def _flush(header_level: Optional[int] = None) -> None:
            nonlocal current_content_blocks
            if not current_content_blocks:
                return
            breadcrumb = self._build_breadcrumb(hierarchy)
            full_text = "\n".join(b["text"] for b in current_content_blocks)
            page_number = current_content_blocks[0].get("page_number")

            if _token_count(full_text) <= max(1, self.config.chunk_size):
                # Section fits in one chunk
                section_text = f"{breadcrumb}\n\n{full_text}".strip() if breadcrumb else full_text
                chunks.append(ParsedChunk(
                    text=section_text,
                    chunk_type=ChunkType.TEXT,
                    chunk_index=len(chunks),
                    page_number=page_number,
                    section_header=breadcrumb or None,
                    parent_header=None,
                    prev_chunk_tail=None,
                    next_chunk_head=None,
                    doc_id=doc_id,
                    source_filename=source_filename,
                ))
            else:
                # Section too large: fallback to fixed chunking with breadcrumb prefix
                fake_block = {
                    "chunk_type": ChunkType.TEXT,
                    "page_number": page_number,
                    "section_header": breadcrumb or None,
                    "parent_header": None,
                }
                sub_chunks = self._chunk_text_block(
                    full_text, fake_block, doc_id, source_filename, len(chunks), ChunkType.TEXT
                )
                # Prepend breadcrumb to each sub-chunk
                for sub in sub_chunks:
                    if breadcrumb:
                        sub.text = f"{breadcrumb}\n\n{sub.text}".strip()
                    sub.chunk_index = len(chunks)
                    chunks.append(sub)

            current_content_blocks = []

        for block in blocks:
            if block["chunk_type"] == ChunkType.HEADER:
                header_level = self._detect_header_level(block["text"], patterns)
                if header_level is not None:
                    _flush(header_level)
                    # Clear sub-levels
                    hierarchy = {k: v for k, v in hierarchy.items() if k < header_level}
                    hierarchy[header_level] = re.sub(r"^#{1,6}\s*", "", block["text"])
                else:
                    # Unrecognized header: treat as content
                    current_content_blocks.append(block)
            else:
                current_content_blocks.append(block)

        _flush()

        # Apply overlap metadata
        for idx, chunk in enumerate(chunks):
            chunk.chunk_index = idx

        return chunks

    # --- Method 4: PCR (Parent-Child Retrieval) ---

    def _chunk_method_pcr(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """PCR: двухуровневая система. Embed child (маленький), retrieve parent (большой)."""
        # Build full text from all non-header blocks
        all_text_parts = []
        for block in blocks:
            if block["chunk_type"] != ChunkType.HEADER and block["text"].strip():
                all_text_parts.append(block["text"])
        full_text = "\n".join(all_text_parts)

        parent_step = max(1, self.config.parent_chunk_size)
        child_step = max(1, self.config.child_chunk_size)
        parent_tokens = _tokenize(full_text)

        chunks: list[ParsedChunk] = []
        parent_idx = 0

        # Шаг Parent: создаем крупные смысловые окна для ответа LLM.
        for p_offset in range(0, len(parent_tokens), parent_step):
            parent_token_slice = parent_tokens[p_offset: p_offset + parent_step]
            parent_text = " ".join(parent_token_slice).strip()
            if not parent_text:
                continue

            parent_id = f"{doc_id}:pcr_parent:{parent_idx}"
            child_tokens = _tokenize(parent_text)

            # Шаг Child: режем parent на мелкие фрагменты для векторного поиска.
            for c_offset in range(0, len(child_tokens), child_step):
                child_token_slice = child_tokens[c_offset: c_offset + child_step]
                child_text = " ".join(child_token_slice).strip()
                if not child_text:
                    continue

                chunks.append(ParsedChunk(
                    text=parent_text,          # Full parent: sent to LLM as context
                    embedding_text=child_text, # Small child: used for precise embedding
                    chunk_type=ChunkType.TEXT,
                    chunk_index=len(chunks),
                    page_number=None,
                    section_header=f"Блок {parent_idx + 1}",
                    parent_header=None,
                    prev_chunk_tail=None,
                    next_chunk_head=None,
                    doc_id=doc_id,
                    source_filename=source_filename,
                    parent_chunk_id=parent_id,
                ))

            parent_idx += 1

        return chunks

    # --- Method 5: Symbol ---

    def _chunk_method_symbol(self, blocks: list[dict], doc_id: str, source_filename: str) -> list[ParsedChunk]:
        """Symbol: разбиение по специальному символу-разделителю, расставленному пользователем."""
        sep = self.config.symbol_separator or "---chunk---"

        # Join all block texts into full document text
        all_text = "\n".join(
            block["text"] for block in blocks
            if block["chunk_type"] != ChunkType.HEADER and block["text"].strip()
        )

        # Пользователь сам управляет семантическими границами через специальный разделитель.
        segments = [seg.strip() for seg in all_text.split(sep) if seg.strip()]

        if not segments:
            # Fallback: treat entire text as one chunk
            segments = [all_text.strip()] if all_text.strip() else []

        chunks: list[ParsedChunk] = []
        for idx, segment in enumerate(segments):
            chunks.append(ParsedChunk(
                text=segment,
                chunk_type=ChunkType.TEXT,
                chunk_index=idx,
                page_number=None,
                section_header=None,
                parent_header=None,
                prev_chunk_tail=None,
                next_chunk_head=None,
                doc_id=doc_id,
                source_filename=source_filename,
            ))

        return chunks

    # --- Utilities ---

    def detect_language(self, text_sample: str) -> str:
        """Определяет язык документа для метаданных; при ошибке возвращает ``unknown``."""
        if not text_sample.strip():
            return "unknown"
        try:
            return detect(text_sample)
        except Exception:
            return "unknown"

    def estimate_chunks_count(self, filepath: str) -> int:
        path = Path(filepath)
        blocks, _ = self._extract_blocks(path, path.suffix.lower())
        total_tokens = sum(_token_count(block["text"]) for block in blocks)
        return max(1, total_tokens // max(1, self.config.chunk_size) + 1)

    def save_parsing_result(self, notebook_id: str, metadata: DocumentMetadata, chunks: list[ParsedChunk]) -> str:
        """Сериализует метаданные и чанки в JSON-файл промежуточного слоя."""
        target_dir = CHUNKS_DIR / notebook_id
        target_dir.mkdir(parents=True, exist_ok=True)
        output = target_dir / f"{metadata.doc_id}.json"
        payload = {
            "metadata": asdict(metadata),
            "chunks": [{**asdict(chunk), "chunk_type": chunk.chunk_type.value} for chunk in chunks],
        }
        output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(output)

    def load_parsing_result(self, notebook_id: str, doc_id: str) -> tuple[DocumentMetadata, list[ParsedChunk]]:
        path = CHUNKS_DIR / notebook_id / f"{doc_id}.json"
        if not path.exists():
            raise FileNotFoundError(path)
        payload = json.loads(path.read_text(encoding="utf-8"))
        metadata = DocumentMetadata(**payload["metadata"])
        chunks = [ParsedChunk(**{**item, "chunk_type": ChunkType(item["chunk_type"])}) for item in payload["chunks"]]
        return metadata, chunks
