"""Экстрактор для DOCX-файлов: параграфы, стили и таблицы."""
# --- Imports ---
from __future__ import annotations

from pathlib import Path
from typing import Optional

from ..models import ChunkType
from .base import BaseExtractor


# --- Models / Classes ---
class DocxExtractor(BaseExtractor):
    """Извлекает DOCX в список блоков; таблицы переводит в markdown-вид."""

    def extract(self, path: Path) -> tuple[list[dict], Optional[int]]:
        """Возвращает (blocks, total_pages)."""
        try:
            from docx import Document
            doc = Document(path)
        except Exception:
            return ([{
                "text": f"Extracted content from {path.name}",
                "chunk_type": ChunkType.TEXT,
                "page_number": None,
                "section_header": None,
                "parent_header": None,
            }], None)
        blocks: list[dict] = []
        current_header: Optional[str] = None
        # Параграфы DOCX конвертируем в блоки с учетом стилей (heading/list/plain).
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style is not None else "").lower()
            if "heading" in style:
                current_header = text
                blocks.append(
                    {
                        "text": text,
                        "chunk_type": ChunkType.HEADER,
                        "page_number": None,
                        "section_header": current_header,
                        "parent_header": None,
                    }
                )
                continue
            if "list" in style:
                marker = "- "
                blocks.append(
                    {
                        "text": f"{marker}{text}",
                        "chunk_type": ChunkType.TEXT,
                        "page_number": None,
                        "section_header": current_header,
                        "parent_header": None,
                    }
                )
                continue
            blocks.append(
                {
                    "text": text,
                    "chunk_type": ChunkType.TEXT,
                    "page_number": None,
                    "section_header": current_header,
                    "parent_header": None,
                }
            )

        # Таблицы приводим к markdown-представлению, чтобы их можно было чанковать как текст.
        for table in doc.tables:
            rows = [[cell.text.strip().replace("|", "\|") for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = rows[0]
            divider = ["---"] * len(header)
            md_lines = [f"| {' | '.join(header)} |", f"| {' | '.join(divider)} |"]
            for row in rows[1:]:
                md_lines.append(f"| {' | '.join(row)} |")
            blocks.append(
                {
                    "text": "\n".join(md_lines),
                    "chunk_type": ChunkType.TABLE,
                    "page_number": None,
                    "section_header": current_header,
                    "parent_header": None,
                }
            )

        return blocks, None
